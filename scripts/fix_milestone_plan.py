"""
Fix the Milestone Execution Plan DOCX to resolve discrepancies identified in the
submission readiness review.

Changes applied:
  1. M1 Task 1: "llamafactory/PubMedQA (10,000 rows)" → "qiaojin/PubMedQA
     pqa_artificial subset (211,269 rows)"
  2. M2 Task 1: "sentence-transformers/all-MiniLM-L6-v2" → "pritamdeka/
     S-PubMedBert-MS-MARCO", "Store 10,000 embeddings" → "Store 209,108
     chunk embeddings"
  3. M2 Task 3: "distilbert-base-uncased" → "dmis-lab/biobert-v1.1"
"""

import sys
import io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from pathlib import Path
import docx

DOC_PATH = Path(__file__).resolve().parents[1] / "docs" / "Milestone_Execution_Plan.docx"
doc = docx.Document(str(DOC_PATH))


def set_para_text(para, new_text):
    """Replace the text of a paragraph while preserving formatting."""
    if para.runs:
        first_run = para.runs[0]
        first_run.text = new_text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(new_text)


def set_cell_text(cell, new_text):
    """Replace the text of a table cell's first paragraph."""
    if cell.paragraphs:
        para = cell.paragraphs[0]
        set_para_text(para, new_text)


# ── Locate relevant paragraphs by searching their text ────────────────────────
changes = []

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()

    # M1 Task 1 — dataset source (llamafactory/PubMedQA → qiaojin/PubMedQA)
    if "llamafactory/PubMedQA" in text:
        old = text
        new = text.replace("llamafactory/PubMedQA", "qiaojin/PubMedQA pqa_artificial subset")
        new = new.replace("10,000 rows", "211,269 rows")
        set_para_text(para, new)
        changes.append(f"Para [{i}]: Fixed dataset source\n  OLD: {old[:80]}...\n  NEW: {new[:80]}...")

    # M2 Task 1 — embedding model (all-MiniLM-L6-v2 → S-PubMedBert-MS-MARCO)
    if "all-MiniLM-L6-v2" in text:
        old = text
        new = text.replace("sentence-transformers/all-MiniLM-L6-v2",
                           "pritamdeka/S-PubMedBert-MS-MARCO")
        new = new.replace("all-MiniLM-L6-v2", "S-PubMedBert-MS-MARCO")
        set_para_text(para, new)
        changes.append(f"Para [{i}]: Fixed embedding model\n  OLD: {old[:80]}...\n  NEW: {new[:80]}...")

    # M2 Task 3 — classifier model (distilbert-base-uncased → dmis-lab/biobert-v1.1)
    if "distilbert-base-uncased" in text:
        old = text
        new = text.replace("distilbert-base-uncased", "dmis-lab/biobert-v1.1")
        set_para_text(para, new)
        changes.append(f"Para [{i}]: Fixed classifier model\n  OLD: {old[:80]}...\n  NEW: {new[:80]}...")

# ── Also fix tables ───────────────────────────────────────────────────────────
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            cell_text = cell.text.strip()

            if "llamafactory/PubMedQA" in cell_text:
                old = cell_text
                new = cell_text.replace("llamafactory/PubMedQA",
                                        "qiaojin/PubMedQA pqa_artificial subset")
                new = new.replace("10,000 rows", "211,269 rows")
                new = new.replace("10,000", "211,269")
                set_cell_text(cell, new)
                changes.append(f"Table {ti}, Row {ri}: Fixed dataset reference")

            if "all-MiniLM-L6-v2" in cell_text:
                new = cell_text.replace("sentence-transformers/all-MiniLM-L6-v2",
                                        "pritamdeka/S-PubMedBert-MS-MARCO")
                new = new.replace("all-MiniLM-L6-v2", "S-PubMedBert-MS-MARCO")
                new = new.replace("10,000 embeddings", "209,108 chunk embeddings")
                set_cell_text(cell, new)
                changes.append(f"Table {ti}, Row {ri}: Fixed embedding model reference")

            if "distilbert-base-uncased" in cell_text:
                new = cell_text.replace("distilbert-base-uncased", "dmis-lab/biobert-v1.1")
                set_cell_text(cell, new)
                changes.append(f"Table {ti}, Row {ri}: Fixed classifier model reference")

# ── Save ─────────────────────────────────────────────────────────────────────
doc.save(str(DOC_PATH))
print(f"✅ Milestone Execution Plan saved → {DOC_PATH}")

# ── Summary ──────────────────────────────────────────────────────────────────
if changes:
    print(f"\n📝 {len(changes)} change(s) applied:")
    for c in changes:
        print(f"  • {c}")
else:
    print("\n⚠️  No target text found — document may already be up-to-date.")
    print("   Searched for: llamafactory/PubMedQA, all-MiniLM-L6-v2, distilbert-base-uncased")
